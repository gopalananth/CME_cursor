#!/usr/bin/env python3
"""
Script to convert AIML_MRRM_Prompts.docx to Markdown format
"""

import sys
import os
from docx import Document

def docx_to_markdown(docx_path, output_path=None):
    """
    Convert a .docx file to Markdown format
    
    Args:
        docx_path (str): Path to the input .docx file
        output_path (str): Path for the output .md file (optional)
    
    Returns:
        str: Path to the generated markdown file
    """
    try:
        # Load the document
        doc = Document(docx_path)
        
        # If no output path specified, create one based on input filename
        if output_path is None:
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            output_path = f"{base_name}.md"
        
        markdown_content = []
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_content.append("")
                continue
            
            # Check for heading styles
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name[-1] if paragraph.style.name[-1].isdigit() else 1
                markdown_content.append(f"{'#' * level} {text}")
            else:
                # Regular paragraph
                markdown_content.append(text)
        
        # Process tables if any
        for table in doc.tables:
            markdown_content.append("")
            # Create markdown table
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    # Header row
                    markdown_content.append("| " + " | ".join(cells) + " |")
                    # Separator row
                    markdown_content.append("| " + " | ".join(["---"] * len(cells)) + " |")
                else:
                    # Data row
                    markdown_content.append("| " + " | ".join(cells) + " |")
            markdown_content.append("")
        
        # Write to markdown file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
        
        print(f"Successfully converted {docx_path} to {output_path}")
        return output_path
        
    except ImportError:
        print("Error: python-docx library not found.")
        print("Please install it using: pip install python-docx")
        return None
    except Exception as e:
        print(f"Error converting document: {e}")
        return None

def main():
    """Main function to handle command line arguments"""
    docx_file = "AIML_MRRM_Prompts.docx"
    
    if not os.path.exists(docx_file):
        print(f"Error: {docx_file} not found in current directory")
        return
    
    # Convert the document
    output_file = docx_to_markdown(docx_file)
    
    if output_file:
        print(f"Conversion complete! Output file: {output_file}")
        print("You can now open and edit the markdown file in your text editor.")

if __name__ == "__main__":
    main()

